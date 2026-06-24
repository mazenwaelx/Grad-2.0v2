"""
File processing module for handling uploaded documents
Supports .docx, .pdf, .xls/.xlsx, .png, .jpg/.jpeg formats
Includes OCR support for scanned PDFs and images via Tesseract
"""
from __future__ import annotations

import io
import os
import re
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any
import hashlib


# LangChain imports
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS

# File processing imports
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
    print("[INFO] OCR libraries loaded successfully (pytesseract + PIL)")
    
    # Configure Tesseract path for Windows if it's not in PATH
    if os.name == 'nt':
        tesseract_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            print(f"[INFO] Tesseract configured: {tesseract_path}")
except Exception as e:
    print(f"[WARNING] OCR not available: {type(e).__name__}: {e}")
    OCR_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

# Gemini Vision for high-accuracy Arabic OCR (preferred over Tesseract)
try:
    import google.generativeai as genai
    _gemini_api_key = os.environ.get("GOOGLE_API_KEY", "").strip()
    if _gemini_api_key:
        genai.configure(api_key=_gemini_api_key)
        GEMINI_VISION_AVAILABLE = True
        print("[INFO] Gemini Vision API configured for high-accuracy OCR")
    else:
        GEMINI_VISION_AVAILABLE = False
        print("[WARNING] GOOGLE_API_KEY not set — Gemini Vision OCR disabled")
except ImportError:
    GEMINI_VISION_AVAILABLE = False
    print("[WARNING] google-generativeai not installed — Gemini Vision OCR disabled")


class FileProcessor:
    """Handles processing of uploaded files and integration with the chat system"""
    
    def __init__(self, embeddings, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the file processor
        
        Args:
            embeddings: Embedding model instance
            chunk_size: Size of text chunks for processing
            chunk_overlap: Overlap between chunks
        """
        self.embeddings = embeddings
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        
        # Store for uploaded file content
        self.uploaded_files: Dict[str, Dict[str, Any]] = {}
        self.file_vectorstores: Dict[str, FAISS] = {}
    
    def get_file_hash(self, file_content: bytes) -> str:
        """Generate a unique hash for file content"""
        return hashlib.md5(file_content, usedforsecurity=False).hexdigest()  # ✅ SECURE: marked as not for security
    
    def is_supported_file(self, filename: str) -> bool:
        """Check if the file type is supported"""
        supported_extensions = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.png', '.jpg', '.jpeg']
        file_ext = Path(filename).suffix.lower()
        return file_ext in supported_extensions
    
    @staticmethod
    def _is_arabic_text_quality_ok(text: str) -> bool:
        """Check if extracted text contains meaningful Arabic content.
        
        PDFs with embedded/custom fonts often extract as garbled text
        (e.g., 'nnnn nnnnnn nnn') — pdfplumber reads the bytes but maps
        them to wrong characters. This detects that situation.
        
        Returns True if text appears to be real Arabic, False if garbled.
        """
        if not text or len(text.strip()) < 20:
            return False
        
        # Count Arabic characters (U+0600 to U+06FF range)
        arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
        # Count total non-whitespace, non-punctuation characters
        alpha_chars = sum(1 for c in text if c.isalpha())
        
        if alpha_chars == 0:
            return False
        
        # For Arabic legal documents, we expect at least 30% Arabic characters
        # among all alphabetic characters. Garbled text will have ~0% Arabic.
        arabic_ratio = arabic_chars / alpha_chars
        
        # Also check for the "nnnn" pattern (common garbled font symptom)
        has_garbled_pattern = bool(re.search(r'(\w)\1{3,}', text))  # 4+ repeated chars
        garbled_n_count = text.count('nnnn')
        
        is_ok = arabic_ratio > 0.3 and garbled_n_count < 3
        
        if not is_ok:
            print(f"[WARNING] Text quality check failed: arabic_ratio={arabic_ratio:.2f}, "
                  f"garbled_n_count={garbled_n_count}, has_garbled_pattern={has_garbled_pattern}")
        
        return is_ok

    def process_pdf(self, file_content: bytes, filename: str) -> List[Document]:
        """Process PDF file and extract text using pdfplumber.
        If extracted text appears garbled (embedded font issue), falls back to OCR.
        """
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing requires pdfplumber. Install with: pip install pdfplumber")
        
        documents = []
        
        try:
            with io.BytesIO(file_content) as pdf_buffer:
                with pdfplumber.open(pdf_buffer) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text and text.strip():
                            doc = Document(
                                page_content=text.strip(),
                                metadata={
                                    "source": filename,
                                    "page": page_num,
                                    "file_type": "pdf",
                                    "uploaded_file": True
                                }
                            )
                            documents.append(doc)
        except Exception as e:
            raise Exception(f"Failed to process PDF '{filename}': {e}")
        
        # Quality check: if pdfplumber extracted text but it's garbled
        # (embedded fonts producing 'nnnn' instead of Arabic), fall back to OCR
        if documents:
            combined_text = " ".join(doc.page_content for doc in documents[:3])
            if not self._is_arabic_text_quality_ok(combined_text):
                print(f"[WARNING] PDF '{filename}' has garbled text (embedded font issue) — falling back to OCR")
                try:
                    ocr_documents = self.process_scanned_pdf(file_content, filename)
                    if ocr_documents:
                        return ocr_documents
                    else:
                        print(f"[WARNING] OCR fallback also produced no text for '{filename}'")
                except Exception as e:
                    print(f"[WARNING] OCR fallback failed for '{filename}': {e}")
        
        return documents

    def _is_scanned_pdf(self, file_content: bytes) -> bool:
        """
        Heuristic: returns True if the PDF has negligible or garbled text,
        indicating it is likely a scanned document that needs OCR.
        Checks the first 3 pages for both text quantity and quality.
        """
        try:
            with io.BytesIO(file_content) as buf:
                import pdfplumber
                with pdfplumber.open(buf) as pdf:
                    sample_pages = pdf.pages[:3]
                    text = " ".join(
                        (page.extract_text() or "") for page in sample_pages
                    )
                    # Check 1: Too little text
                    if len(text.strip()) < 50:
                        return True
                    # Check 2: Text exists but is garbled (embedded fonts)
                    if not self._is_arabic_text_quality_ok(text):
                        print(f"[INFO] PDF has garbled text — treating as scanned")
                        return True
                    return False
        except Exception:
            return True  # If we can't open it normally, assume scanned

    def process_scanned_pdf(self, file_content: bytes, filename: str) -> List[Document]:
        """
        Process a scanned PDF by converting pages to images then extracting text.
        Uses Gemini Vision (primary) or Tesseract OCR (fallback) for text extraction.
        """
        if not PDF2IMAGE_AVAILABLE:
            raise ImportError(
                "Scanned PDF processing requires pdf2image. "
                "Install with: pip install pdf2image (and install Poppler)"
            )
        if not GEMINI_VISION_AVAILABLE and not OCR_AVAILABLE:
            raise ImportError(
                "OCR requires either google-generativeai (Gemini Vision) or pytesseract. "
                "Install one of them."
            )

        documents = []

        try:
            # Convert PDF pages to images at high DPI for better OCR
            print(f"[INFO] Converting scanned PDF pages to images: {filename}")
            images = convert_from_bytes(file_content, dpi=300)

            for page_num, page_image in enumerate(images, 1):
                # Convert PIL image to bytes for Gemini Vision
                img_buffer = io.BytesIO()
                page_image.save(img_buffer, format='PNG')
                page_bytes = img_buffer.getvalue()
                
                # Try Gemini Vision first (much better for Arabic + numbers)
                text = None
                if GEMINI_VISION_AVAILABLE:
                    text = self._extract_text_with_gemini_vision(page_bytes)
                
                # Fallback to Tesseract if Gemini failed or unavailable
                if not text and OCR_AVAILABLE:
                    print(f"[INFO] Falling back to Tesseract OCR for page {page_num}")
                    processed_image = self._preprocess_image_for_ocr(page_image)
                    text = self._extract_text_ordered(processed_image)

                if text and text.strip():
                    ocr_method = "gemini_vision" if GEMINI_VISION_AVAILABLE else "tesseract"
                    doc = Document(
                        page_content=text.strip(),
                        metadata={
                            "source": filename,
                            "page": page_num,
                            "file_type": "pdf",
                            "ocr": True,
                            "ocr_method": ocr_method,
                            "uploaded_file": True,
                        },
                    )
                    documents.append(doc)

        except Exception as e:
            raise Exception(f"Failed to OCR scanned PDF '{filename}': {e}")

        return documents

    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image to improve OCR accuracy for Arabic legal documents.
        Uses gentle preprocessing to preserve number/digit details.
        Applies: upscaling, grayscale, moderate contrast, noise reduction, Otsu binarization.
        """
        try:
            from PIL import ImageEnhance, ImageFilter
            import numpy as np
            
            # Step 1: Upscale small images — Tesseract works best at ~300 DPI
            # For a typical A4 page at 300 DPI, width should be ~2480px
            MIN_WIDTH = 2000
            w, h = image.size
            if w < MIN_WIDTH:
                scale = MIN_WIDTH / w
                new_size = (int(w * scale), int(h * scale))
                image = image.resize(new_size, Image.LANCZOS)
                print(f"[INFO] Upscaled image from {w}x{h} to {new_size[0]}x{new_size[1]}")
            
            # Step 2: Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Step 3: Moderate contrast enhancement (1.5 instead of 2.0 to preserve numbers)
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)
            
            # Step 4: Light sharpening to bring out text edges
            image = image.filter(ImageFilter.SHARPEN)
            
            # Step 5: Gentle noise reduction (size=3 median filter)
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            # Step 6: Otsu's binarization — proper implementation
            # This automatically finds the optimal threshold instead of using mean,
            # which preserves subtle features like number strokes.
            img_array = np.array(image)
            
            # Otsu's method: find threshold that minimizes intra-class variance
            hist, _ = np.histogram(img_array, bins=256, range=(0, 256))
            total_pixels = img_array.size
            
            best_threshold = 0
            best_variance = 0
            
            weight_bg = 0
            sum_bg = 0
            total_sum = sum(i * hist[i] for i in range(256))
            
            for t in range(256):
                weight_bg += hist[t]
                if weight_bg == 0:
                    continue
                weight_fg = total_pixels - weight_bg
                if weight_fg == 0:
                    break
                
                sum_bg += t * hist[t]
                mean_bg = sum_bg / weight_bg
                mean_fg = (total_sum - sum_bg) / weight_fg
                
                variance = weight_bg * weight_fg * (mean_bg - mean_fg) ** 2
                if variance > best_variance:
                    best_variance = variance
                    best_threshold = t
            
            img_array = np.where(img_array > best_threshold, 255, 0).astype(np.uint8)
            image = Image.fromarray(img_array)
            
            print(f"[INFO] Image preprocessing completed (Otsu threshold={best_threshold})")
            return image
            
        except Exception as e:
            print(f"[WARNING] Image preprocessing failed: {e}, using original image")
            return image
    
    def _extract_text_ordered(self, image: Image.Image) -> str:
        """
        Extract text from image using Tesseract with position-aware ordering.
        Uses image_to_data() to get bounding boxes, then sorts text blocks
        top-to-bottom to maintain the correct reading order of the document.
        Falls back to image_to_string() if position extraction fails.
        """
        try:
            import pandas as pd
            
            # Use image_to_data to get text with position information
            # PSM 3 = Automatic page segmentation (detects blocks, but may order them wrong)
            # OEM 3 = LSTM + Legacy engine
            custom_config = r'--oem 3 --psm 3'
            
            data = pytesseract.image_to_data(
                image,
                lang='ara+eng',
                config=custom_config,
                output_type=pytesseract.Output.DATAFRAME
            )
            
            # Filter out empty/low-confidence entries
            # conf = -1 means no text detected for that entry  
            data = data[data['conf'] != -1].copy()
            
            if data.empty:
                print(f"[WARNING] image_to_data returned no text, falling back to image_to_string")
                return self._extract_text_simple(image)
            
            # Group text by block number — each block_num is a detected text region
            # Sort blocks by their top-Y coordinate (top of page first)
            blocks = []
            for block_num in data['block_num'].unique():
                block_data = data[data['block_num'] == block_num]
                
                # Get the top-Y position of this block (minimum top value)
                block_top = block_data['top'].min()
                block_left = block_data['left'].min()
                
                # Reconstruct text within this block: group by line, then by word order
                lines = []
                for line_num in sorted(block_data['line_num'].unique()):
                    line_data = block_data[block_data['line_num'] == line_num]
                    # Sort words by their left position (handles both LTR and RTL)
                    line_data = line_data.sort_values('left')
                    words = line_data['text'].astype(str).tolist()
                    line_text = ' '.join(w for w in words if w.strip())
                    if line_text.strip():
                        lines.append(line_text.strip())
                
                block_text = '\n'.join(lines)
                if block_text.strip():
                    blocks.append((block_top, block_left, block_text.strip()))
            
            if not blocks:
                print(f"[WARNING] No text blocks reconstructed, falling back to image_to_string")
                return self._extract_text_simple(image)
            
            # Sort blocks by Y position (top to bottom) — this fixes the ordering issue
            # For blocks at similar Y positions (same row), sort right-to-left for Arabic
            LINE_TOLERANCE = 30  # pixels — blocks within this range are on the same "row"
            blocks.sort(key=lambda b: (b[0] // LINE_TOLERANCE, -b[1]))
            
            # Join all blocks with double newline
            ordered_text = '\n\n'.join(block[2] for block in blocks)
            
            print(f"[INFO] Extracted {len(blocks)} text blocks in correct top-to-bottom order")
            return ordered_text
            
        except Exception as e:
            print(f"[WARNING] Position-aware extraction failed: {e}, falling back to image_to_string")
            return self._extract_text_simple(image)
    
    def _extract_text_simple(self, image: Image.Image) -> str:
        """
        Simple fallback: extract text using image_to_string with PSM 6.
        PSM 6 = Assume a single uniform block of text — reads strictly top-to-bottom.
        """
        custom_config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(
            image,
            lang='ara+eng',
            config=custom_config
        )
        return text if text else ""

    def _extract_text_with_gemini_vision(self, image_bytes: bytes) -> str:
        """
        Extract text from an image using Gemini Vision API.
        Far more accurate than Tesseract for Arabic text, numbers, and dates.
        Returns the extracted text or empty string on failure.
        """
        try:
            import google.generativeai as genai
            
            # Try multiple Gemini models for vision (preview models get deprecated)
            model = None
            model_candidates = [
                'gemini-2.5-flash',           # Stable 2.5 Flash
                'gemini-2.0-flash',           # Stable 2.0 Flash
                'gemini-1.5-flash',           # Fallback
            ]
            for model_name in model_candidates:
                try:
                    model = genai.GenerativeModel(model_name)
                    break
                except Exception:
                    continue
            
            if model is None:
                print(f"[WARNING] No Gemini Vision model available")
                return ""
            
            # Detect MIME type from image bytes (don't hardcode!)
            if image_bytes[:3] == b'\xff\xd8\xff':
                mime_type = "image/jpeg"
            elif image_bytes[:4] == b'\x89PNG':
                mime_type = "image/png"
            else:
                mime_type = "image/png"  # Default fallback
            
            # Create image part for the API
            image_part = {
                "mime_type": mime_type,
                "data": image_bytes
            }
            
            # Carefully crafted prompt for verbatim Arabic text extraction
            prompt = """أنت أداة استخراج نص متخصصة. استخرج كل النص الموجود في هذه الصورة بالضبط كما هو مكتوب.

القواعد:
1. انسخ النص حرفياً كما يظهر في الصورة - لا تغير أي كلمة أو رقم أو تاريخ
2. حافظ على الترتيب من أعلى الصفحة إلى أسفلها
3. انسخ جميع الأرقام والتواريخ وأرقام المواد القانونية وأرقام الأحكام بدقة تامة
4. لا تضف أي شرح أو تعليق أو تلخيص - فقط انسخ النص الموجود
5. حافظ على فواصل الأسطر والفقرات كما في الأصل
6. إذا كان هناك نص غير واضح، اكتبه كما تراه بأفضل تقدير

استخرج النص الآن:"""
            
            # Call Gemini Vision API
            response = model.generate_content(
                [prompt, image_part],
                generation_config=genai.types.GenerationConfig(
                    temperature=0.0,  # Zero temperature for exact extraction
                    max_output_tokens=8192,
                )
            )
            
            if response and response.text:
                extracted_text = response.text.strip()
                print(f"[INFO] Gemini Vision extracted {len(extracted_text)} chars")
                print(f"[DEBUG] Gemini Vision preview: {extracted_text[:200]}")
                return extracted_text
            else:
                # Retry once with simplified prompt (handles edge cases)
                print(f"[WARNING] Gemini Vision returned empty — retrying with simplified prompt")
                response = model.generate_content(
                    ["Extract all Arabic text from this image exactly as written:", image_part],
                    generation_config=genai.types.GenerationConfig(
                        temperature=0.1,
                        max_output_tokens=8192,
                    )
                )
                if response and response.text:
                    extracted_text = response.text.strip()
                    print(f"[INFO] Gemini Vision retry extracted {len(extracted_text)} chars")
                    return extracted_text
                
                print(f"[WARNING] Gemini Vision returned empty response after retry")
                return ""
                
        except Exception as e:
            print(f"[WARNING] Gemini Vision extraction failed: {type(e).__name__}: {e}")
            return ""

    def process_image(self, file_content: bytes, filename: str) -> List[Document]:
        """
        Process an image file (PNG/JPG) for text extraction.
        Uses Gemini Vision API (primary) for high-accuracy Arabic OCR,
        with Tesseract as fallback if Gemini is unavailable.
        """
        if not GEMINI_VISION_AVAILABLE and not OCR_AVAILABLE:
            raise ImportError(
                "Image processing requires either google-generativeai (Gemini Vision) "
                "or pytesseract + Pillow. Install one of them."
            )

        documents = []
        file_ext = Path(filename).suffix.lower()

        try:
            print(f"[INFO] Processing image: {filename}")
            
            # Strategy: Try Gemini Vision first (much better for Arabic + numbers)
            text = None
            ocr_method = "unknown"
            
            if GEMINI_VISION_AVAILABLE:
                print(f"[INFO] Using Gemini Vision for high-accuracy OCR")
                text = self._extract_text_with_gemini_vision(file_content)
                if text and text.strip():
                    ocr_method = "gemini_vision"
            
            # Fallback to Tesseract if Gemini failed or unavailable
            if (not text or not text.strip()) and OCR_AVAILABLE:
                print(f"[INFO] Falling back to Tesseract OCR for {filename}")
                image = Image.open(io.BytesIO(file_content))
                print(f"[INFO] Image size: {image.size}")
                processed_image = self._preprocess_image_for_ocr(image)
                text = self._extract_text_ordered(processed_image)
                if text and text.strip():
                    ocr_method = "tesseract"
            
            print(f"[DEBUG] OCR ({ocr_method}) extracted {len(text) if text else 0} chars from {filename}")
            if text and text.strip():
                print(f"[DEBUG] OCR text preview: {text.strip()[:200]}")
                
                doc = Document(
                    page_content=text.strip(),
                    metadata={
                        "source": filename,
                        "file_type": file_ext.lstrip("."),
                        "ocr": True,
                        "ocr_method": ocr_method,
                        "uploaded_file": True,
                    },
                )
                documents.append(doc)
            else:
                print(f"[WARNING] No text extracted from {filename}")
                doc = Document(
                    page_content=f"صورة مرفوعة: {filename} - لم يتم استخراج نص من هذه الصورة. قد تكون الصورة لا تحتوي على نص قابل للقراءة.",
                    metadata={
                        "source": filename,
                        "file_type": file_ext.lstrip("."),
                        "ocr": True,
                        "uploaded_file": True,
                        "empty_ocr": True,
                    },
                )
                documents.append(doc)

        except Exception as e:
            raise Exception(f"Failed to process image '{filename}': {e}")

        return documents
    
    def process_docx(self, file_content: bytes, filename: str) -> List[Document]:
        """Process DOCX file and extract text"""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing requires python-docx. Install with: pip install python-docx")
        
        documents = []
        
        try:
            with io.BytesIO(file_content) as docx_buffer:
                doc = DocxDocument(docx_buffer)
                
                # Extract paragraphs
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text.strip())
                
                # Extract tables
                tables_text = []
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_data.append(cell.text.strip())
                        if row_data:
                            table_data.append(" | ".join(row_data))
                    if table_data:
                        tables_text.append("\n".join(table_data))
                
                # Combine all text
                all_text = "\n\n".join(paragraphs)
                if tables_text:
                    all_text += "\n\n" + "\n\n".join(tables_text)
                
                if all_text.strip():
                    doc = Document(
                        page_content=all_text.strip(),
                        metadata={
                            "source": filename,
                            "file_type": "docx",
                            "uploaded_file": True
                        }
                    )
                    documents.append(doc)
        
        except Exception as e:
            raise Exception(f"Failed to process DOCX file: {e}")
        
        return documents
    
    def process_excel(self, file_content: bytes, filename: str) -> List[Document]:
        """Process Excel file and extract text"""
        if not PANDAS_AVAILABLE:
            raise ImportError("Excel processing requires pandas and openpyxl. Install with: pip install pandas openpyxl")
        
        documents = []
        
        try:
            with io.BytesIO(file_content) as excel_buffer:
                # Read all sheets
                excel_file = pd.ExcelFile(excel_buffer)
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(excel_buffer, sheet_name=sheet_name)
                    
                    # Convert DataFrame to text
                    if not df.empty:
                        # Create a readable text representation
                        text_lines = [f"Sheet: {sheet_name}"]
                        text_lines.append("=" * 50)
                        
                        # Add column headers
                        headers = " | ".join(str(col) for col in df.columns)
                        text_lines.append(headers)
                        text_lines.append("-" * len(headers))
                        
                        # Add rows
                        for _, row in df.iterrows():
                            row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row)
                            text_lines.append(row_text)
                        
                        sheet_text = "\n".join(text_lines)
                        
                        doc = Document(
                            page_content=sheet_text,
                            metadata={
                                "source": filename,
                                "sheet": sheet_name,
                                "file_type": "excel",
                                "uploaded_file": True
                            }
                        )
                        documents.append(doc)
        
        except Exception as e:
            raise Exception(f"Failed to process Excel file: {e}")
        
        return documents
    
    def process_file(self, file_content: bytes, filename: str) -> Tuple[List[Document], str]:
        """
        Process uploaded file and return documents
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (documents, file_hash)
        """
        file_ext = Path(filename).suffix.lower()
        file_hash = self.get_file_hash(file_content)
        
        # Check if already processed
        if file_hash in self.uploaded_files:
            return self.uploaded_files[file_hash]["documents"], file_hash
        
        documents = []
        
        try:
            if file_ext in ['.png', '.jpg', '.jpeg']:
                documents = self.process_image(file_content, filename)
            elif file_ext == '.pdf':
                # Detect scanned PDFs and route to OCR pipeline
                if self._is_scanned_pdf(file_content):
                    print(f"[INFO] Detected scanned PDF: {filename} — using OCR")
                    documents = self.process_scanned_pdf(file_content, filename)
                else:
                    documents = self.process_pdf(file_content, filename)
            elif file_ext in ['.docx', '.doc']:
                documents = self.process_docx(file_content, filename)
            elif file_ext in ['.xlsx', '.xls']:
                documents = self.process_excel(file_content, filename)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Split documents into chunks
            # For OCR files (images/scanned PDFs), skip splitting if text is small
            # enough to keep as a single document — preserves all details and avoids
            # slow embedding on CPU
            is_ocr_file = any(doc.metadata.get("ocr", False) for doc in documents)
            total_text_len = sum(len(doc.page_content) for doc in documents)
            
            if is_ocr_file and total_text_len < 8000:
                # Keep OCR text as-is — no splitting needed for small documents
                chunked_documents = documents
                print(f"[INFO] OCR file ({total_text_len} chars) — skipping text splitting")
            else:
                chunked_documents = []
                for doc in documents:
                    chunks = self.text_splitter.split_documents([doc])
                    chunked_documents.extend(chunks)
                
                # If no chunks were created but we had documents, use originals
                if not chunked_documents and documents:
                    print(f"[WARNING] Text splitter produced no chunks, using original documents")
                    chunked_documents = documents
            
            print(f"[INFO] File {filename}: {len(documents)} docs -> {len(chunked_documents)} chunks")
            
            # Store file information
            self.uploaded_files[file_hash] = {
                "filename": filename,
                "documents": chunked_documents,
                "original_documents": documents,
                "file_type": file_ext,
                "is_ocr": is_ocr_file,
                "processed_at": pd.Timestamp.now() if PANDAS_AVAILABLE else "unknown"
            }
            
            return chunked_documents, file_hash
            
        except Exception as e:
            raise Exception(f"Error processing file {filename}: {str(e)}")
    
    def create_file_vectorstore(self, file_hash: str) -> FAISS:
        """Create a FAISS vectorstore for the uploaded file"""
        if file_hash not in self.uploaded_files:
            raise ValueError(f"File with hash {file_hash} not found")
        
        if file_hash in self.file_vectorstores:
            return self.file_vectorstores[file_hash]
        
        documents = self.uploaded_files[file_hash]["documents"]
        
        if not documents:
            raise ValueError("No documents found for the file")
        
        # Create vectorstore
        vectorstore = FAISS.from_documents(documents, self.embeddings)
        self.file_vectorstores[file_hash] = vectorstore
        
        return vectorstore
    
    def search_in_file(self, file_hash: str, query: str, k: int = 3) -> List[Document]:
        """Search for relevant content in the uploaded file.
        For OCR files (images/scanned PDFs), returns ALL content directly
        instead of doing vector search — faster and preserves all details.
        """
        if file_hash not in self.uploaded_files:
            raise ValueError(f"File with hash {file_hash} not found")
        
        file_info = self.uploaded_files[file_hash]
        
        # For OCR files, return ALL documents directly — no vector search needed
        # This is faster (skips embedding) and preserves all content
        if file_info.get("is_ocr", False):
            print(f"[INFO] OCR file — returning full content directly (no vector search)")
            return file_info["documents"]
        
        # For non-OCR files (DOCX, Excel, large PDFs), use vector search
        if file_hash not in self.file_vectorstores:
            self.create_file_vectorstore(file_hash)
        
        vectorstore = self.file_vectorstores[file_hash]
        retriever = vectorstore.as_retriever(search_kwargs={"k": k})
        
        return retriever.invoke(query)
    
    def get_file_info(self, file_hash: str) -> Optional[Dict[str, Any]]:
        """Get information about an uploaded file"""
        return self.uploaded_files.get(file_hash)
    
    def remove_file(self, file_hash: str) -> bool:
        """Remove an uploaded file from memory"""
        removed = False
        
        if file_hash in self.uploaded_files:
            del self.uploaded_files[file_hash]
            removed = True
        
        if file_hash in self.file_vectorstores:
            del self.file_vectorstores[file_hash]
            removed = True
        
        return removed
    
    def list_uploaded_files(self) -> List[Dict[str, Any]]:
        """List all uploaded files"""
        files_info = []
        for file_hash, file_data in self.uploaded_files.items():
            files_info.append({
                "hash": file_hash,
                "filename": file_data["filename"],
                "file_type": file_data["file_type"],
                "document_count": len(file_data["documents"]),
                "processed_at": file_data.get("processed_at", "unknown")
            })
        return files_info
    
    def calculate_question_relevance(self, question: str, file_hash: str) -> float:
        """
        Calculate how relevant a question is to the uploaded file content
        
        Returns a score between 0.0 and 1.0:
        - 1.0: Explicitly about the file
        - 0.7-0.9: High relevance based on content similarity
        - 0.3-0.6: Medium relevance
        - 0.0-0.2: Low relevance
        """
        if file_hash not in self.uploaded_files:
            return 0.0
        
        file_info = self.uploaded_files[file_hash]
        filename = file_info["filename"]
        question_lower = question.lower()
        
        # Explicit file keywords (highest priority)
        explicit_keywords = [
            "الملف", "المستند", "الوثيقة", "المرفق", "الملف المرفوع",
            "في الملف", "من الملف", "حسب الملف", "وفقا للملف",
            "المحتوى", "النص", "البيانات", "المعلومات المرفقة"
        ]
        
        for keyword in explicit_keywords:
            if keyword in question_lower:
                return 1.0  # Explicitly about the file
        
        # Check filename mention
        filename_base = Path(filename).stem.lower()
        if len(filename_base) > 3 and filename_base in question_lower:
            return 0.9  # Very high relevance
        
        # Content-based relevance scoring
        try:
            # Get a sample of the file content for comparison
            documents = file_info["documents"][:5]  # Check first 5 chunks
            file_content = " ".join([doc.page_content.lower() for doc in documents])
            
            # Extract meaningful words from question (remove stop words)
            stop_words = {
                'في', 'من', 'إلى', 'على', 'عن', 'مع', 'هذا', 'هذه', 'التي', 'الذي', 
                'ما', 'هل', 'كيف', 'متى', 'أين', 'لماذا', 'كم', 'أي', 'كل', 'بعض',
                'قد', 'لقد', 'كان', 'كانت', 'يكون', 'تكون', 'أن', 'إن', 'لكن', 'لكن',
                'أو', 'أم', 'بل', 'لا', 'لن', 'لم', 'ما', 'غير', 'سوى', 'إلا'
            }
            
            question_words = set(question_lower.split()) - stop_words
            question_words = {word for word in question_words if len(word) > 2}  # Remove very short words
            
            if not question_words:
                return 0.1  # No meaningful words to compare
            
            # Calculate word overlap
            file_words = set(file_content.split())
            overlap = len(question_words.intersection(file_words))
            overlap_ratio = overlap / len(question_words) if question_words else 0
            
            # Check for specific names or entities mentioned in the file
            file_entities = set()
            # Extract potential names (Arabic names pattern)
            import re
            arabic_names = re.findall(r'[أ-ي]{3,}(?:\s+[أ-ي]{3,})*', file_content)
            for name in arabic_names:
                if len(name.split()) <= 3:  # Reasonable name length
                    file_entities.add(name.lower())
            
            # Check if question mentions specific entities from the file
            entity_overlap = 0
            for entity in file_entities:
                if entity in question_lower:
                    entity_overlap += 1
            
            # Boost score for specific domain terms
            contract_terms = {
                'عقد', 'اتفاقية', 'شروط', 'بنود', 'راتب', 'أجر', 'مكافأة', 'مدة',
                'موظف', 'عامل', 'شركة', 'طرف', 'التزامات', 'مسؤوليات', 'ساعات'
            }
            
            legal_terms = {
                'مادة', 'فقرة', 'قانون', 'نظام', 'حقوق', 'واجبات', 'تشريع',
                'إجازة', 'تأمين', 'تعويض', 'فصل', 'استقالة'
            }
            
            contract_overlap = len(question_words.intersection(contract_terms))
            legal_overlap = len(question_words.intersection(legal_terms))
            
            # Calculate boosts
            entity_boost = min(entity_overlap * 0.3, 0.6)  # Strong boost for entity mentions
            contract_boost = min(contract_overlap * 0.2, 0.4)  # Up to 0.4 boost for contract terms
            legal_penalty = min(legal_overlap * 0.1, 0.3)  # Penalty for legal terms (suggests general legal question)
            
            # Calculate final relevance score
            base_score = overlap_ratio * 0.5  # Base score from word overlap
            final_score = base_score + entity_boost + contract_boost - legal_penalty
            final_score = max(0.0, min(final_score, 1.0))  # Clamp between 0 and 1
            
            # Additional penalty for very general legal questions
            general_legal_terms = {'قانون', 'حقوق', 'واجبات', 'مادة', 'نظام', 'تشريع'}
            if len(question_words.intersection(general_legal_terms)) >= 2:
                final_score *= 0.3  # Strong penalty for general legal questions
            
            return final_score
            
        except Exception as e:
            print(f"[ERROR] Failed to calculate relevance for file {filename}: {e}")
            return 0.1  # Low relevance on error
    
    def is_question_about_file(self, question: str, file_hash: str) -> bool:
        """
        Determine if a question is about the uploaded file content
        Uses the relevance score with a threshold and additional logic
        """
        if file_hash not in self.uploaded_files:
            return False
        
        question_lower = question.lower()
        
        # Explicit file mentions - always about the file
        explicit_file_keywords = [
            "الملف", "المستند", "الوثيقة", "المرفق", "الملف المرفوع",
            "في الملف", "من الملف", "حسب الملف", "وفقا للملف",
            # Image-related keywords
            "الصورة", "صورة", "اشرح الصورة", "الصوره", "اشرح لي الصورة",
            "اللي رفعته", "رفعته", "المرفوعة", "المرفوع",
            "محتوى الصورة", "في الصورة", "من الصورة",
        ]
        
        for keyword in explicit_file_keywords:
            if keyword in question_lower:
                return True
        
        # If the file was processed via OCR (image/scanned PDF), 
        # treat questions as being about the file by default
        file_info = self.uploaded_files.get(file_hash, {})
        file_docs = file_info.get("documents", [])
        is_ocr_file = any(
            doc.metadata.get("ocr", False) for doc in file_docs
        ) if file_docs else False
        
        if is_ocr_file:
            # For OCR files, only switch away if the question is explicitly about the legal database
            strong_legal_indicators_count = sum(
                1 for indicator in ["قانون العمل المصري", "حسب القانون", "وفقا للقانون", "القانون المصري"]
                if indicator in question_lower
            )
            if strong_legal_indicators_count == 0:
                return True
        
        # Strong legal indicators - likely NOT about the file
        strong_legal_indicators = [
            "قانون العمل المصري", "حسب القانون", "وفقا للقانون", "القانون المصري",
            "النظام", "التشريع"
        ]
        
        # Weaker legal indicators that could be in both file and legal database
        weak_legal_indicators = [
            "المادة", "حقوق العامل", "واجبات العامل", "العواقب القانونية"
        ]
        
        # Count indicators
        strong_legal_count = sum(1 for indicator in strong_legal_indicators if indicator in question_lower)
        weak_legal_count = sum(1 for indicator in weak_legal_indicators if indicator in question_lower)
        
        # If question has 2+ strong legal indicators, it's probably about the legal database
        if strong_legal_count >= 2:
            return False
        
        # If question has 1 strong + 2 weak indicators, likely legal database
        if strong_legal_count >= 1 and weak_legal_count >= 2:
            return False
        
        # Check for specific article numbers (likely legal database)
        import re
        if re.search(r'المادة\s*\d+', question_lower):
            return False
        
        # Use relevance score for borderline cases
        relevance = self.calculate_question_relevance(question, file_hash)
        
        # Check for contract-specific terms that suggest file relevance
        contract_indicators = ['طرف', 'عقد', 'اتفاقية', 'شروط', 'بنود']
        has_contract_terms = any(term in question_lower for term in contract_indicators)
        
        # Check for workplace-related terms that could be in files
        workplace_indicators = ['بيئة عمل', 'مكان العمل', 'السلامة المهنية', 'وسائل السلامة', 'صاحب العمل']
        has_workplace_terms = any(term in question_lower for term in workplace_indicators)
        
        # Check for procedure/process questions that are likely legal
        procedure_indicators = ['كيف يتم', 'كيفية', 'إجراءات', 'خطوات']
        has_procedure_terms = any(term in question_lower for term in procedure_indicators)
        
        # Determine threshold based on question characteristics
        total_legal_count = strong_legal_count + weak_legal_count
        
        if (total_legal_count >= 2 and not has_contract_terms and not has_workplace_terms) or has_procedure_terms:
            return relevance >= 0.6  # Higher threshold for clearly legal questions
        elif has_contract_terms or has_workplace_terms:
            return relevance >= 0.15  # Lower threshold for contract/workplace questions
        elif total_legal_count >= 1:
            return relevance >= 0.45  # Medium threshold for questions with some legal terms
        else:
            return relevance >= 0.35  # Normal threshold for other questions


# Global file processor instance (will be initialized in the UI)
_file_processor: Optional[FileProcessor] = None

def get_file_processor() -> Optional[FileProcessor]:
    """Get the global file processor instance"""
    return _file_processor

def set_file_processor(processor: FileProcessor):
    """Set the global file processor instance"""
    global _file_processor
    _file_processor = processor