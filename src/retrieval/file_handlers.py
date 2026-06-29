"""
File-type handler implementations (Strategy pattern).

Each handler knows how to extract ``Document`` objects from
a specific file format.
"""
from __future__ import annotations

import io
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import List

from langchain_core.documents import Document

from src.retrieval.ocr_strategies import CompositeOCR

# ── Optional dependencies ──────────────────────────────────────────

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    pd = None  # type: ignore[assignment]
    PANDAS_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None  # type: ignore[assignment,misc]
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    pdfplumber = None  # type: ignore[assignment]
    PDF_AVAILABLE = False

try:
    from pdf2image import convert_from_bytes
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    convert_from_bytes = None  # type: ignore[assignment]
    PDF2IMAGE_AVAILABLE = False


# ── Abstract base ──────────────────────────────────────────────────

class FileTypeHandler(ABC):
    """Abstract base for processing a specific file type."""

    @abstractmethod
    def process(self, file_content: bytes, filename: str) -> List[Document]:
        """Extract documents from raw file content."""

    @staticmethod
    def _is_arabic_text_quality_ok(text: str) -> bool:
        if not text or len(text.strip()) < 20:
            return False
        arabic = sum(1 for c in text if "\u0600" <= c <= "\u06FF")
        alpha = sum(1 for c in text if c.isalpha())
        if alpha == 0:
            return False
        return arabic / alpha > 0.3 and text.count("nnnn") < 3


# ── PDF ────────────────────────────────────────────────────────────

class PDFHandler(FileTypeHandler):
    """Handle PDF files — text extraction with OCR fallback."""

    def __init__(self, ocr: CompositeOCR) -> None:
        self._ocr = ocr

    def process(self, file_content: bytes, filename: str) -> List[Document]:
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber required — pip install pdfplumber")

        docs = self._extract_text(file_content, filename)

        # Quality check → OCR fallback
        if docs:
            combined = " ".join(d.page_content for d in docs[:3])
            if not self._is_arabic_text_quality_ok(combined):
                ocr_docs = self._ocr_all_pages(file_content, filename)
                if ocr_docs:
                    return ocr_docs

        if not docs and self.is_scanned(file_content):
            return self._ocr_all_pages(file_content, filename)

        return docs

    def is_scanned(self, file_content: bytes) -> bool:
        try:
            with io.BytesIO(file_content) as buf, pdfplumber.open(buf) as pdf:
                text = " ".join((p.extract_text() or "") for p in pdf.pages[:3])
                return len(text.strip()) < 50 or not self._is_arabic_text_quality_ok(text)
        except Exception:
            return True

    def _extract_text(self, content: bytes, filename: str) -> List[Document]:
        docs = []
        with io.BytesIO(content) as buf, pdfplumber.open(buf) as pdf:
            for num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    docs.append(Document(
                        page_content=text.strip(),
                        metadata={"source": filename, "page": num, "file_type": "pdf", "uploaded_file": True},
                    ))
        return docs

    def _ocr_all_pages(self, content: bytes, filename: str) -> List[Document]:
        if not PDF2IMAGE_AVAILABLE:
            raise ImportError("pdf2image required — pip install pdf2image")
        if not self._ocr.is_available:
            raise ImportError("OCR engine required (google-generativeai or pytesseract)")

        docs = []
        images = convert_from_bytes(content, dpi=300)
        for num, img in enumerate(images, 1):
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            text = self._ocr.extract_text(buf.getvalue())
            if text and text.strip():
                docs.append(Document(
                    page_content=text.strip(),
                    metadata={"source": filename, "page": num, "file_type": "pdf", "ocr": True, "uploaded_file": True},
                ))
        return docs


# ── DOCX ───────────────────────────────────────────────────────────

class DocxHandler(FileTypeHandler):
    def process(self, file_content: bytes, filename: str) -> List[Document]:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required — pip install python-docx")

        with io.BytesIO(file_content) as buf:
            doc = DocxDocument(buf)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            tables_text = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    tables_text.append("\n".join(rows))

            all_text = "\n\n".join(paragraphs)
            if tables_text:
                all_text += "\n\n" + "\n\n".join(tables_text)

            if all_text.strip():
                return [Document(
                    page_content=all_text.strip(),
                    metadata={"source": filename, "file_type": "docx", "uploaded_file": True},
                )]
        return []


# ── Excel ──────────────────────────────────────────────────────────

class ExcelHandler(FileTypeHandler):
    def process(self, file_content: bytes, filename: str) -> List[Document]:
        if not PANDAS_AVAILABLE:
            raise ImportError("pandas + openpyxl required — pip install pandas openpyxl")

        docs = []
        with io.BytesIO(file_content) as buf:
            xls = pd.ExcelFile(buf)
            for sheet in xls.sheet_names:
                df = pd.read_excel(buf, sheet_name=sheet)
                if df.empty:
                    continue
                lines = [f"Sheet: {sheet}", "=" * 50]
                lines.append(" | ".join(str(c) for c in df.columns))
                for _, row in df.iterrows():
                    lines.append(" | ".join(str(v) if pd.notna(v) else "" for v in row))
                docs.append(Document(
                    page_content="\n".join(lines),
                    metadata={"source": filename, "sheet": sheet, "file_type": "excel", "uploaded_file": True},
                ))
        return docs


# ── Image ──────────────────────────────────────────────────────────

class ImageHandler(FileTypeHandler):
    def __init__(self, ocr: CompositeOCR) -> None:
        self._ocr = ocr

    def process(self, file_content: bytes, filename: str) -> List[Document]:
        if not self._ocr.is_available:
            raise ImportError("OCR engine required (google-generativeai or pytesseract)")

        ext = Path(filename).suffix.lower().lstrip(".")
        text = self._ocr.extract_text(file_content)

        if text and text.strip():
            return [Document(
                page_content=text.strip(),
                metadata={"source": filename, "file_type": ext, "ocr": True, "uploaded_file": True},
            )]

        return [Document(
            page_content=f"صورة مرفوعة: {filename} - لم يتم استخراج نص من هذه الصورة.",
            metadata={"source": filename, "file_type": ext, "ocr": True, "uploaded_file": True, "empty_ocr": True},
        )]
